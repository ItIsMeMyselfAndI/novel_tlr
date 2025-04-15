from docx import Document
import json
import os
import requests
import time
import winsound


def displayFooter():
    print("="*50)
    print("[*] Wait for a few moments ...")
    print("="*50)
    time.sleep(3)


def getAISpecs(specs):
    role = f"You are a professional novel translator of {specs["in_lang"]} to {specs["out_lang"]} language."
    messages = [
        {"role":"system", "content":role},
        {"role":"user", "content":"I want you to translate something using the prompt i gave you earlier."},
        {"role":"assistant", "content":""},
        {"role":"user", "content":""}
    ]
    return messages


def getFileNames(raw_folder):
    files = os.listdir(raw_folder)
    file_names = []
    for ch in files[:]:
        if ("~$" not in ch) and (ch != ".git"):
            file_names.append(ch)
    return file_names


def getChaps(raw_folder, file_names):
    chaps = {}
    for ch in file_names:
        i = ch.index(".")
        document = Document(f"{raw_folder}\\{ch}")
        content = "\n".join([p.text for p in document.paragraphs])
        chaps[ch[:i]] = content
    return chaps


def getChapterDictionary(ch_dictionary_file, ch):
    with open(ch_dictionary_file, "r", encoding="utf-8") as file:
        text = file.read().strip()
    all_dictionary = text.split("\n\n---\n\n")
    for ch_dictionary in all_dictionary:
        if f"Chapter {ch}" in ch_dictionary:
            return ch_dictionary
    return "{None}"


def getTranslation(ch, messages, retries):
    url="https://openrouter.ai/api/v1/chat/completions"
    key = "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" # ds v3 0324
    # key = "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" 
    # key = "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"
    # key = "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"
    # key = "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"
    model="deepseek/deepseek-chat:free"

    for _ in range(retries + 1):
        response = requests.post(
            url=url,
            headers={
                "Authorization": f"Bearer {key}",
                "Content-Type": "application/json",
            },
            data=json.dumps({
                "model": model,
                "messages": messages
            })
        )
        response_json = response.json()
        usage = response_json.get("usage", {})
        if "choices" in response_json.keys():
            chap_translation = response_json["choices"][0]["message"]["content"]
            if "**CHAPTER END**" in chap_translation:
                print(f"[*] Successful translation - {ch}")
                winsound.Beep(1000, 1*1000) # hz, ms
                return chap_translation, usage
            print(f"[!] Error: Incomplete extraction - {ch} >> Restarting ...")

        else:
            print(f"[!] Error: No response - {ch} >> Restarting ...")
        winsound.Beep(1000, int(1.5*1000)) # hz, ms
    return None, usage 


def saveChap(text, folder, file_name):
    document = Document()
    document.add_heading(file_name)
    document.add_paragraph(text)
    document.save(f"{folder}\\{file_name}.docx")


def displayUsageInfo(usage):
    # completion = usage.get("completion_tokens_details", {})
    info = (
        f"{"":>4}>> Prompt: {usage.get("prompt_tokens"):<8}"
        f"Output: {usage.get("completion_tokens", 0):<8}"
        # f"\t\tReasoning tokens: {completion.get("reasoning_tokens", 0)}\n"
        # f"\t\tAccepted prediction tokens: {completion.get("accepted_prediction_tokens", 0)}\n"
        # f"\t\tRejected Prediction tokens: {completion.get("rejected_prediction_tokens", 0)}\n"
        f"Total: {usage.get("total_tokens")}"
    )
    print(info)


def main():
    title = "SL"

    series_folder = f".\\books\\{title}"
    raw_folder = os.path.join(series_folder, "raw")
    tl_folder = os.path.join(series_folder, "tl")
    
    os.makedirs(tl_folder, exist_ok=True)
    # context_file = os.path.join(series_folder, "context.txt")
    ch_dictionary_file = os.path.join(series_folder, "ch_dictionary.txt")
    
    with open("translator_instruction.txt", "r", encoding="utf-8") as file:
        instruction = file.read()
    
    with open("translator_confirmation.txt", "r", encoding="utf-8") as file:
        confirmation = f"Understood. To recap, I will follow this instruction: {instruction}\n"
        confirmation += file.read()
    
    specs = {
        "in_lang":"{hangul}", 
        "out_lang":"{english}", 
        "retries": 20,
        "instruction": instruction
    }
    messages = getAISpecs(specs)
    
    file_names = getFileNames(raw_folder)
    if not file_names:
        os.system("cls")
        print("="*50)
        print("[!] Error: Folder containing the raw chapters is empty")
        print("[*] Exiting the program ...")
        print("="*50)
        time.sleep(3)
        return
    chaps = getChaps(raw_folder, file_names)

    # with open(context_file, "r", encoding="utf-8") as file:
    #     context = file.read().strip()

    for i, ch in enumerate(chaps.keys()):
        ch_dictionary = getChapterDictionary(ch_dictionary_file, ch)
        if i != 0:
            print("\n" + "-"*50 + "\n")
        print(ch_dictionary)
        print("\n" + "-"*50 + "\n")
        
        confirmation += f"\nAnd cross-check other keywords based from this dictionary:\n{ch_dictionary}."
        messages[2]["content"] = confirmation
        messages[3]["content"] = f"translate this:\n{chaps[ch]}"
        chap_translation, usage = getTranslation(ch, messages, specs["retries"])
        if chap_translation in ["", None]:
            print(f"[!] Unsuccessful extraction - {ch}")
            continue
        saveChap(chap_translation, tl_folder, ch)
        displayUsageInfo(usage)
        time.sleep(0.5)
  
    # alarm
    for i in range(3):
        winsound.Beep(1000, int(0.5*1000))
        time.sleep(0.3)
    return tl_folder

            
if __name__ == "__main__":
    os.system("cls")
    prev = time.time()
    print("="*50)
    print("[*] Results")
    print("="*50)
    print()

    out_path = main()
    
    curr = time.time()
    elapsed = curr - prev
    print("="*50)
    print(f"[*] Elapsed time is {elapsed}s")
    print(f"[*] Output path is '{out_path}'")
    print("="*50)
