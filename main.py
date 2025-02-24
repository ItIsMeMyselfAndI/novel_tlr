import json
import os
import requests
import sys
import time
from docx import Document


def getAISpecs():
    messages = []
    in_lang, out_lang = _getLanguages()
    instruction = _getInstruction()
    context = _getContext()
    messages.append(
        {"role": "system", "content": f"you are a professional novel translator from {in_lang} to {out_lang}"}
    )
    messages.append(
        {"role": "assistant", 
        "content": f"Understood. I will follow this instruction: {instruction}\nAnd I will base my translation from this context: {context}"}
    )
    return messages

def _getLanguages():
    os.system("cls")
    print("="*50)
    in_lang = input("[*] Enter input Language: ")
    out_lang = input("[*] Enter output Language: ")
    print("="*50)
    print("[*] Wait for a few moments ...")
    print("="*50)
    time.sleep(3)
    return in_lang, out_lang

def _getInstruction():
    os.system("cls")
    print("="*50)
    print("[*] Enter the instruction to guide the AI.")
    print("[*] Press {Enter} >> {Ctrl + Z} >> {Enter} in order once done.")
    print("="*50)
    instruction = sys.stdin.read().strip()
    print("="*50)
    print("[*] Wait for a few moments ...")
    print("="*50)
    time.sleep(3)
    return instruction

def _getContext():
    os.system("cls")
    print("="*50)
    print("[*] Enter the context/settings for the AI to base upon.")
    print("[*] Press {Enter} >> {Ctrl + Z} >> {Enter} in order once done.")
    print("="*50)
    context = sys.stdin.read().strip()
    print("="*50)
    print("[*] Wait for a few moments ...")
    print("="*50)
    time.sleep(3)
    return context

# def getValidChapRange():
#     os.system("cls")
#     print("="*50)
#     while True:
#         try:
#             chap_range = input("Enter chapter range (e.g. 10-15): ").strip().split("-")
#             chap_range = [int(float(chap_range[0])), int(float(chap_range[1]))]
#         except IndexError:
#             print("[!] Error: Only enter lower and upper bounds.")
#             continue
#         except ValueError:
#             print("[!] Error: Only enter integers.")
#             continue
#         print("="*50)
#         print("[*] Wait for a few moments ...")
#         print("="*50)
#         time.sleep(3)
#         return chap_range


def getFolderPaths():
    os.system("cls")
    print("="*50)
    while True:
        print("[*] Enter folder path of the raw chapter(s):")
        raw_folder_path = os.path.abspath(input("\t>> "))
        if  not os.path.isdir(raw_folder_path):
            print("[!] Error: Folder path does not exist.")
            continue
        print("="*50)
        break
    print("[*] Enter folder path of the translated chapter(s):")
    tl_folder_path = os.path.abspath(input("\t>> "))
    os.makedirs(tl_folder_path, exist_ok=True)
    print("="*50)
    print("[*] Wait for a few moments ...")
    print("="*50)
    time.sleep(3)
    return raw_folder_path, tl_folder_path


def getFileNames(raw_folder_path):
    files = os.listdir(raw_folder_path)
    file_names = []
    for name in files[:]:
        if ("~$" not in name) and (name != ".git"):
            file_names.append(name)
    return file_names


def getChaps(raw_folder_path, file_names):
    chaps = {}
    for name in file_names:
        i = name.index(".")
        document = Document(f"{raw_folder_path}\\{name}")
        content = "\n".join([p.text for p in document.paragraphs])
        chaps[name[:i]] = content
    return chaps


def getNumOfRetries():
    print("="*50)
    while True:
        try:
            print("[*] Enter number of re-tries for failed translation of each chapter.")
            print("[*] Enter '0' if you don't want to re-try.")
            retries = int(float(input("\t>> ")))
        except ValueError:
            print("[!] Error: Invalid input. Only enter an integer.")
            continue
        print("="*50)
        print("[*] Wait for a few moments ...")
        print("="*50)
        time.sleep(3)
        return retries


def getTranslation(url, key, model, messages, content):
    os.system("cls")
    messages_cp = messages[:]
    messages_cp.append({"role": "user", "content" : f"translate this:\n{content}"})
    response = requests.post(
        url=url,
        headers={
            "Authorization": f"Bearer {key}",
            "Content-Type": "application/json",
        },
        data=json.dumps({
            "model": model,
            "messages": messages_cp
        })
    )
    response_json = response.json()
    if "choices" in response_json.keys():
        chap_translation = response_json["choices"][0]["message"]["content"]
        return chap_translation


def saveChap(text, folder_path, file_name):
    document = Document()
    document.add_heading(file_name)
    document.add_paragraph(text)
    document.save(f"{folder_path}\\{file_name}.docx")


def main():
    url="https://openrouter.ai/api/v1/chat/completions"
    # key = "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"
    # model = "deepseek/deepseek-r1:free"
    key = "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"
    model="deepseek/deepseek-chat:free"

    raw_folder_path, tl_folder_path = getFolderPaths()
    messages = getAISpecs()
    # chap_range = getValidChapRange()
    # chaps = getChaps(chap_range)
    file_names = getFileNames(raw_folder_path)
    chaps = getChaps(raw_folder_path, file_names)
    retries = getNumOfRetries()

    prev = time.time()
    os.system("cls")
    
    print("="*50)
    print("[*] Results")
    print("="*50)
   
    for name in chaps.keys():
        chap_translation = getTranslation(url, key, model, messages, chaps[name])
        for i in range(retries + 1):
            if (chap_translation not in ["", None]) and ("**CHAPTER END**" in chap_translation):
                saveChap(chap_translation, tl_folder_path, name)
                print(f"[*] Successful translation - {name}")
                break
            elif (chap_translation in ["", None]):
                print(f"[!] Error: No response - {name} >> Restarting translation ...")
            else:
                print(f"[!] Error: Incomplete translation - {name} >> Restarting translation ...")
            chap_translation = getTranslation(url, key, model, messages, chaps[name])
  
    curr = time.time()
    elapsed = curr - prev

    print("="*50)
    print(f"[*] Elapsed time is {elapsed}")
    print(f"[*] Output path is '{tl_folder_path}'")
    print("="*50)
            

if __name__ == "__main__":
    main()