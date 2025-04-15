from docx import Document
import os
import time


def rplWordInTLChs(tl_folder, old, new):
    try:
        tl_files = [f for f in os.listdir(tl_folder) if (".git" not in f) and ("~" not in f)]
    except FileNotFoundError:
        print("[!] Folder was not found")
        return

    chs = {}
    for f in tl_files:
        exist = False
        path = os.path.join(tl_folder, f)
        doc = Document(path)
        paragraphs = []
        # print(len(doc.paragraphs))
        for p in doc.paragraphs[1:]:
            # print(p.text)
            # print("="*50)
            if old in p.text:
                exist = True
            paragraphs.append(p.text)
        
        if exist:
            chs[f] = paragraphs

    if len(chs) == 0:
        print("[*] Old string(s) not found")
        return
    # return

    for f_name, content in chs.items():
        print(f"[*] Replacing string(s) >> {f_name}")
        path = os.path.join(tl_folder, f_name)
        doc = Document()
        doc.add_heading(f_name.split(".")[0])
        for p in content:
            doc.add_paragraph(p.replace(old, new))
        doc.save(path)

        

if __name__ == "__main__":
    tl_folder = ".\\books\\SL\\tl"
    # new = "Endurance"
    # old = "Tenacity"
    old = "Middle-rank"
    new = "Mid-rank"

    os.system("cls")
    prev = time.time()
    print("="*50)
    print(f"[*] Replacing {{{old}}} >> {{{new}}} ...")
    print("="*50)

    rplWordInTLChs(tl_folder, old, new)

    print("="*50)
    curr = time.time()
    elapsed = curr - prev
    print(f"[*] Elapsed time is {elapsed}s")
    print(f"[*] Output path is '{tl_folder}'")
    print("="*50)
    